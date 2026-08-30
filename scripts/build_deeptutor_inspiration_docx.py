#!/usr/bin/env python3
"""Build the DeepTutor inspiration strategy brief as a polished DOCX.

The report body is authored in Markdown and rendered here with an explicit
standard_business_brief token map plus an editorial_cover first page.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "outputs" / "deeptutor-inspiration-analysis-2026-08-23.md"
OUTPUT_DOCX = ROOT / "outputs" / "DeepTutor对本项目的启发与落地路线图_2026-08-23.docx"
ASSET_DIR = ROOT / "outputs" / "assets"
ARCH_DIAGRAM = ASSET_DIR / "deeptutor-target-architecture.png"

# standard_business_brief preset with a CJK-safe named font override.  Using
# one Unicode family for every OOXML font slot prevents headless renderers from
# classifying Chinese runs as hAnsi and substituting a non-CJK fallback.
BASE_FONT = "Arial Unicode MS"
CJK_FONT = "Arial Unicode MS"
BODY_SIZE = 11
BODY_AFTER = 6
BODY_LINE = 1.10
H1_BLUE = "2E74B5"
H2_BLUE = "2E74B5"
H3_BLUE = "1F4D78"
NAVY = "203748"
TEAL = "2B5163"
GOLD = "B07A2A"
MUTED = "5B6570"
LIGHT_FILL = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FBF4E8"
TABLE_BORDER = "C7D0D9"
WHITE = "FFFFFF"
BLACK = "111111"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120

INLINE_TOKEN = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\)|<https?://[^>]+>|\[[DPL]\d+\])"
)

_BOOKMARK_ID = 1000


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run,
    *,
    name: str = BASE_FONT,
    cjk: str = CJK_FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), cjk)
    rfonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color: str = BLACK, bold: bool = False) -> None:
    style.font.name = BASE_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), BASE_FONT)
    rfonts.set(qn("w:hAnsi"), BASE_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rfonts.set(qn("w:cs"), BASE_FONT)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=BODY_SIZE)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(BODY_AFTER)
    pf.line_spacing = BODY_LINE
    pf.widow_control = True

    h1 = styles["Heading 1"]
    set_style_font(h1, size=16, color=H1_BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True
    h1.paragraph_format.page_break_before = False

    h2 = styles["Heading 2"]
    set_style_font(h2, size=13, color=H2_BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = styles["Heading 3"]
    set_style_font(h3, size=12, color=H3_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    for style_name in ("Caption",):
        style = styles[style_name]
        set_style_font(style, size=9, color=MUTED)
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)

    if "Source Note" not in styles:
        source = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source = styles["Source Note"]
    set_style_font(source, size=8.5, color=MUTED)
    source.font.italic = True
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)
    source.paragraph_format.line_spacing = 1.0

    if "List Body Exact" not in styles:
        ls = styles.add_style("List Body Exact", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ls = styles["List Body Exact"]
    set_style_font(ls, size=BODY_SIZE)
    ls.paragraph_format.space_before = Pt(0)
    ls.paragraph_format.space_after = Pt(8)
    ls.paragraph_format.line_spacing = 1.167
    ls.paragraph_format.widow_control = True

    if "Table Body" not in styles:
        ts = styles.add_style("Table Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ts = styles["Table Body"]
    set_style_font(ts, size=8.8)
    ts.paragraph_format.space_before = Pt(0)
    ts.paragraph_format.space_after = Pt(2)
    ts.paragraph_format.line_spacing = 1.0


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=TABLE_BORDER, size="5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa=TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_hyperlink(paragraph, text: str, url: str, *, color=H1_BLUE, underline=True):
    part = paragraph.part
    rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BASE_FONT)
    rfonts.set(qn("w:hAnsi"), BASE_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rpr.append(rfonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_internal_hyperlink(paragraph, text: str, anchor: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BASE_FONT)
    rfonts.set(qn("w:hAnsi"), BASE_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), H1_BLUE)
    rpr.append(color)
    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    rpr.append(vert)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmarked_source_marker(paragraph, marker: str) -> None:
    global _BOOKMARK_ID
    bookmark_id = str(_BOOKMARK_ID)
    _BOOKMARK_ID += 1
    anchor = f"src_{marker[1:-1]}"
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), anchor)
    paragraph._p.append(start)
    run = paragraph.add_run(marker)
    set_run_font(run, size=BODY_SIZE, color=H1_BLUE, bold=True)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.append(end)


def add_inline(paragraph, text: str, *, base_size: float | None = None, base_color: str = BLACK) -> None:
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Menlo", cjk=CJK_FONT, size=(base_size or BODY_SIZE) - 0.7, color=TEAL)
            rpr = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF2F4")
            rpr.append(shd)
        elif token.startswith("["):
            md_link = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if md_link:
                label, url = md_link.groups()
                add_hyperlink(paragraph, label, url)
            else:
                add_internal_hyperlink(paragraph, token, f"src_{token[1:-1]}")
        elif token.startswith("<"):
            url = token[1:-1]
            add_hyperlink(paragraph, url, url)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=base_color)


def add_field_run(paragraph, field: str, *, size=9, color=MUTED):
    run = paragraph.add_run()
    set_run_font(run, size=size, color=color)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def set_page_furniture(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("AICOURSE.TOP  ·  产品与研究策略")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=2)
    run = p.add_run("\tDEEPTUTOR EVIDENCE BRIEF")
    set_run_font(run, size=8.5, color=MUTED)

    first_header = section.first_page_header
    fp = first_header.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run("AICOURSE.TOP  ·  RESEARCH & STRATEGY")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    for footer in (section.footer, section.first_page_footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("第 ")
        set_run_font(run, size=8.5, color=MUTED)
        add_field_run(p, "PAGE", size=8.5)
        run = p.add_run(" 页  ·  DeepTutor 对本项目的启发与落地路线图")
        set_run_font(run, size=8.5, color=MUTED)


def paragraph_shading_border(paragraph, *, fill: str, border_color: str, border_size="18") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    p_bdr = ppr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        ppr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), border_size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("证据型产品研究简报")
    set_run_font(run, size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("DeepTutor 对本项目的\n启发与落地路线图")
    set_run_font(run, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("从“证据优先的静态课程平台”走向“可审计的闭环学习系统”")
    set_run_font(run, size=14.5, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("aicourse.top / Agentic Engineering 项目")
    set_run_font(run, size=11, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("快照 2026-08-23  ·  DeepTutor v1.5.16  ·  commit 8515dfd…")
    set_run_font(run, size=9.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("生产部署状态：NOT_ASSESSABLE  ·  本报告只核验当前工作树")
    set_run_font(run, size=9.5, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.18
    paragraph_shading_border(p, fill=PALE_BLUE, border_color=H1_BLUE)
    run = p.add_run("核心判断  ")
    set_run_font(run, size=10.5, color=NAVY, bold=True)
    add_inline(
        p,
        "DeepTutor 最值得借鉴的不是功能数量，也不是在每页加一个聊天框，而是把课程材料、证据检索、对话辅导、学习者证据、针对性练习和效果评估连成可追溯、可纠错的闭环。",
        base_size=10.5,
        base_color=NAVY,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("建议路线：保留静态课程为权威发布面，增加独立、可选、可撤销的 Tutor 服务面")
    set_run_font(run, size=10, color=GOLD, bold=True)
    doc.add_page_break()


def add_reading_map(doc: Document) -> None:
    p = doc.add_paragraph("内容导航", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    items = [
        ("判断", "执行摘要、范围与证据边界"),
        ("对照", "本项目基线、DeepTutor 架构与证据强度"),
        ("设计", "课程内容、AI Tutor、RAG、Memory 与其他工程启发"),
        ("落地", "双平面目标架构、路线图与 P0/P1/P2 优先级"),
        ("验证", "成功指标、真实学习实验与风险清单"),
        ("证据", "固定版本来源、本地定位与复核说明"),
    ]
    for label, desc in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(f"{label}  ")
        set_run_font(run, size=10.5, color=H1_BLUE, bold=True)
        run = p.add_run(desc)
        set_run_font(run, size=10.5, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    paragraph_shading_border(p, fill=PALE_GOLD, border_color=GOLD, border_size="12")
    run = p.add_run("阅读提示  ")
    set_run_font(run, size=9.5, color=NAVY, bold=True)
    add_inline(
        p,
        "“已核实事实”“合理推断”“本项目建议”在报告中保持分离；当前工作树不等于生产上线，论文代理指标也不等于真实学习增益。",
        base_size=9.5,
        base_color=NAVY,
    )


def add_numbering_defs(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abs_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abs_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def make_abstract(abs_id: int, fmt: str, text: str, bullet_font: str | None = None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
        lvl.append(ppr)
        if bullet_font:
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), bullet_font)
            rfonts.set(qn("w:hAnsi"), bullet_font)
            rpr.append(rfonts)
            lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)

    make_abstract(next_abs, "bullet", "•", "Arial")
    bullet_abs = next_abs
    make_abstract(next_abs + 1, "decimal", "%1.")
    number_abs = next_abs + 1

    def make_num(num_id: int, abs_id: int):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)

    make_num(next_num, bullet_abs)
    make_num(next_num + 1, number_abs)
    return next_num, next_num + 1


def new_number_instance(doc: Document, abstract_num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    next_num = max((int(x.get(qn("w:numId"))) for x in nums), default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abs_ref)
    numbering.append(num)
    return next_num


def get_abstract_num_id(doc: Document, num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    for num in numbering.findall(qn("w:num")):
        if int(num.get(qn("w:numId"))) == num_id:
            return int(num.find(qn("w:abstractNumId")).get(qn("w:val")))
    raise KeyError(num_id)


def apply_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, nid])


def choose_table_widths(headers: list[str]) -> list[int]:
    joined = "|".join(headers)
    if headers and headers[0] == "模块":
        return [1450, 2300, 2350, 3260]
    if headers and headers[0] == "阶段":
        return [1350, 3300, 2850, 1860]
    if headers and headers[0] == "状态":
        return [1550, 3900, 3910]
    if headers and headers[0] == "维度":
        return [1050, 1900, 1900, 1900, 2610]
    if headers and headers[0] == "风险":
        return [1600, 2500, 2800, 2460]
    count = len(headers)
    if count == 2:
        return [2700, 6660]
    if count == 3:
        return [1900, 3530, 3930]
    if count == 4:
        return [1800, 2200, 2480, 2880]
    base = CONTENT_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    headers, body = rows[0], rows[1:]
    table_size = 8.0 if len(headers) >= 5 else 8.6
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, choose_table_widths(headers))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Body"]
        p.paragraph_format.space_after = Pt(1)
        add_inline(p, text, base_size=table_size, base_color=NAVY)
        for run in p.runs:
            run.bold = True
    for row_data in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cell = cells[idx]
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Body"]
            add_inline(p, text, base_size=table_size)
    set_table_geometry(table, choose_table_widths(headers))
    note = doc.add_paragraph(style="Source Note")
    note.add_run("来源：DeepTutor 固定提交与论文、本项目当前工作树，以及本报告明确标注的设计建议；详见附录 A。")


def draw_centered_multiline(draw: ImageDraw.ImageDraw, box, text: str, font, fill, spacing=10) -> None:
    x0, y0, x1, y1 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((x0 + x1 - width) / 2, (y0 + y1 - height) / 2), text, font=font, fill=fill, spacing=spacing, align="center")


def make_architecture_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1800, 980
    img = Image.new("RGB", (w, h), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    font_path = "/System/Library/Fonts/STHeiti Medium.ttc"
    light_path = "/System/Library/Fonts/STHeiti Light.ttc"
    title_font = ImageFont.truetype(font_path, 52)
    label_font = ImageFont.truetype(font_path, 31)
    body_font = ImageFont.truetype(light_path, 25)
    small_font = ImageFont.truetype(light_path, 22)

    draw.text((90, 58), "推荐目标架构：双平面 + 课程学习契约", font=title_font, fill=f"#{NAVY}")
    draw.text((92, 128), "静态课程保持权威与低运维；Tutor 作为独立、可选、可撤销的服务运行。", font=body_font, fill=f"#{MUTED}")

    boxes = {
        "left": (85, 245, 535, 595),
        "center": (675, 245, 1125, 595),
        "right": (1265, 245, 1715, 595),
        "bottom": (265, 735, 1535, 900),
    }
    fills = {"left": "EAF2F8", "center": "FBF4E8", "right": "E8F3F1", "bottom": "F2F4F7"}
    strokes = {"left": H1_BLUE, "center": GOLD, "right": TEAL, "bottom": MUTED}
    for key, box in boxes.items():
        draw.rounded_rectangle(box, radius=28, fill=f"#{fills[key]}", outline=f"#{strokes[key]}", width=5)

    draw_centered_multiline(draw, boxes["left"], "静态发布面\n\n课程 · manifest · 来源\n确定性互动 · 本地进度\n\n权威版本 / 免费 / 无需账户", label_font, f"#{NAVY}", spacing=12)
    draw_centered_multiline(draw, boxes["center"], "课程学习契约\n\n目标 · 概念 · 先修 · 锚点\n权限 · content hash\n索引 / 题库 / 评测版本", label_font, f"#{NAVY}", spacing=12)
    draw_centered_multiline(draw, boxes["right"], "Tutor 服务面\n\nExplain · Coach · Practice\nRAG · 学习者证据 · 工具\n会话 · 同意 · 导出 / 删除", label_font, f"#{NAVY}", spacing=12)
    draw_centered_multiline(draw, boxes["bottom"], "观测与评估层\n检索  |  引文  |  教学动作  |  学习行为  |  安全  |  延迟与成本\n只保存可审计事件与结果证据，不保存隐藏思维链", label_font, f"#{NAVY}", spacing=10)

    def arrow(start, end, color):
        draw.line([start, end], fill=f"#{color}", width=8)
        ex, ey = end
        sx, sy = start
        import math

        angle = math.atan2(ey - sy, ex - sx)
        size = 22
        pts = [
            (ex, ey),
            (ex - size * math.cos(angle - 0.55), ey - size * math.sin(angle - 0.55)),
            (ex - size * math.cos(angle + 0.55), ey - size * math.sin(angle + 0.55)),
        ]
        draw.polygon(pts, fill=f"#{color}")

    arrow((545, 420), (660, 420), H1_BLUE)
    arrow((1135, 420), (1250, 420), TEAL)
    arrow((310, 610), (495, 720), MUTED)
    arrow((900, 610), (900, 720), MUTED)
    arrow((1490, 610), (1305, 720), MUTED)
    draw.text((520, 377), "发布", font=small_font, fill=f"#{MUTED}")
    draw.text((1137, 377), "约束", font=small_font, fill=f"#{MUTED}")
    img.save(path, format="PNG", optimize=True)


def set_picture_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_architecture_figure(doc: Document) -> None:
    make_architecture_diagram(ARCH_DIAGRAM)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    shape = run.add_picture(str(ARCH_DIAGRAM), width=Inches(6.35))
    set_picture_alt_text(
        shape,
        "双平面目标架构",
        "左侧为静态课程发布面，中间为版本化课程学习契约，右侧为可选 Tutor 服务面，三者共同向下连接观测与评估层。",
    )
    caption = doc.add_paragraph(style="Caption")
    caption.add_run("图 1　推荐目标架构：静态发布面通过课程学习契约约束 Tutor 服务面，并共享可审计的观测与评估层。")


def markdown_body_lines() -> list[str]:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## 执行摘要")
    return lines[start:]


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw.append(lines[i].strip())
        i += 1
    rows: list[list[str]] = []
    for idx, line in enumerate(raw):
        parts = [x.strip() for x in line.strip("|").split("|")]
        if idx == 1 and all(re.fullmatch(r":?-{3,}:?", p) for p in parts):
            continue
        rows.append(parts)
    return rows, i


PAGE_BREAK_HEADINGS = {
    "2. 本项目当前基线：优势与结构性缺口",
    "5. 课程内容：建议新增“AI Tutor 与学习系统工程”",
    "6. AI Tutor：从功能清单变成闭环协议",
    "7. RAG：从教学示例到课程级证据服务",
    "9. 推荐目标架构：双平面＋课程学习契约",
    "10. 分阶段路线图",
    "12. 风险、缓解与停止条件",
    "附录 A：主要证据来源",
}


def render_markdown(doc: Document) -> None:
    lines = markdown_body_lines()
    bullet_num, default_number_num = add_numbering_defs(doc)
    number_abs = get_abstract_num_id(doc, default_number_num)
    ordered_num = default_number_num
    in_ordered = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            in_ordered = False
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_paragraph(style="Heading 1")
            if text in PAGE_BREAK_HEADINGS and len(doc.paragraphs) > 5:
                # Attach the break to the heading itself.  A standalone break
                # paragraph can be pushed onto an otherwise empty page when
                # the preceding content already fills its page.
                p.paragraph_format.page_break_before = True
            add_inline(p, text, base_size=16, base_color=H1_BLUE)
            if text == "9. 推荐目标架构：双平面＋课程学习契约":
                add_architecture_figure(doc)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:].strip(), base_size=13, base_color=H2_BLUE)
            i += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:].strip(), base_size=12, base_color=H3_BLUE)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines):
            rows, i = parse_table(lines, i)
            if rows:
                add_markdown_table(doc, rows)
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.right_indent = Inches(0.06)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.14
            paragraph_shading_border(p, fill=PALE_BLUE, border_color=H1_BLUE, border_size="14")
            add_inline(p, stripped[2:], base_size=10.5, base_color=NAVY)
            i += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Body Exact")
            apply_num(p, bullet_num)
            bullet_text = bullet_match.group(1)
            source_match = re.match(r"^(\[[DPL]\d+\])\s*(.*)$", bullet_text)
            if source_match:
                add_bookmarked_source_marker(p, source_match.group(1))
                if source_match.group(2):
                    p.add_run(" ")
                    add_inline(p, source_match.group(2), base_size=BODY_SIZE)
            else:
                add_inline(p, bullet_text, base_size=BODY_SIZE)
            i += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number_match:
            if not in_ordered:
                ordered_num = new_number_instance(doc, number_abs)
                in_ordered = True
            p = doc.add_paragraph(style="List Body Exact")
            apply_num(p, ordered_num)
            add_inline(p, number_match.group(1), base_size=BODY_SIZE)
            i += 1
            continue

        # Join wrapped prose lines until the next block marker.
        parts = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or re.match(r"^(#{2,4}\s|[-]\s|\d+\.\s|>\s|\|)", nxt):
                break
            parts.append(nxt)
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        add_inline(p, " ".join(parts), base_size=BODY_SIZE)


def set_document_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "DeepTutor 对本项目的启发与落地路线图"
    props.subject = "课程内容、AI Tutor、RAG、学习者记忆、评估与分阶段实施建议"
    props.author = "Codex"
    props.keywords = "DeepTutor, AI Tutor, RAG, 教育技术, 课程设计, 学习者模型, aicourse.top"
    props.comments = "Evidence-grounded strategy brief; snapshot 2026-08-23."


def add_update_fields_setting(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build() -> Path:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(SOURCE_MD)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_properties(doc)
    configure_styles(doc)
    for section in doc.sections:
        set_page_furniture(section)
    add_update_fields_setting(doc)
    add_cover(doc)
    add_reading_map(doc)
    render_markdown(doc)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = build()
    print(path)
